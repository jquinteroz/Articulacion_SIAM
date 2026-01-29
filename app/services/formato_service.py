"""
Servicio para generación de formatos con datos del aprendiz y acudiente
"""
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyPDF2 import PdfMerger
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
import logging

logger = logging.getLogger(__name__)

def convert_docx_to_pdf(docx_path):
    """
    Convierte un archivo DOCX a PDF usando docx2pdf

    Args:
        docx_path: Ruta del archivo DOCX

    Returns:
        str: Ruta del archivo PDF generado
    """
    try:
        from docx2pdf import convert

        # Generar nombre del PDF (mismo nombre pero con extensión .pdf)
        pdf_path = docx_path.replace('.docx', '.pdf')

        # Convertir
        convert(docx_path, pdf_path)

        # Eliminar el archivo DOCX original
        if os.path.exists(docx_path):
            os.remove(docx_path)

        return pdf_path

    except ImportError:
        logger.warning("docx2pdf no está instalado, retornando DOCX sin convertir")
        return docx_path
    except Exception as e:
        logger.error(f"Error al convertir DOCX a PDF: {str(e)}, retornando DOCX")
        return docx_path

def remove_all_highlights(doc):
    """
    Elimina todos los resaltados amarillos del documento completo
    Optimizado para procesar una sola vez
    """
    # Eliminar highlights en párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.highlight_color = None

    # Eliminar highlights en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = None

def replace_text_in_paragraph(paragraph, replacements):
    """
    Reemplaza texto en un párrafo de forma optimizada
    Procesa todos los reemplazos en una sola pasada
    """
    full_text = paragraph.text

    # Verificar si hay algo que reemplazar
    needs_replacement = any(key in full_text for key in replacements.keys())

    if not needs_replacement:
        return

    # Hacer todos los reemplazos en el texto completo
    for key, value in replacements.items():
        full_text = full_text.replace(key, str(value))

    # Si el texto cambió, actualizar el párrafo
    if full_text != paragraph.text:
        runs = paragraph.runs
        if runs:
            # Limpiar todos los runs excepto el primero
            for run in runs[1:]:
                run.text = ""
            # Poner todo el texto en el primer run
            runs[0].text = full_text

def generar_formato_tratamiento_datos(aprendiz):
    """
    Genera el formato de Tratamiento de Datos con los campos resaltados autocompletados

    Args:
        aprendiz: Objeto Aprendiz con los datos del estudiante y acudiente

    Returns:
        str: Ruta del archivo generado
    """
    try:
        # Cargar la plantilla base
        template_path = os.path.join('formatos', 'Tratamiento de Datos.docx')
        doc = Document(template_path)

        # Datos del aprendiz (convertir None a string vacío)
        nombre_aprendiz = str(aprendiz.usuario.nombre_completo or "") if aprendiz and aprendiz.usuario else ""
        documento_aprendiz = str(aprendiz.usuario.documento or "") if aprendiz and aprendiz.usuario else ""
        tipo_documento_aprendiz = str(aprendiz.usuario.tipo_documento or "") if aprendiz and aprendiz.usuario else ""

        # Datos del acudiente (están en el modelo Aprendiz - convertir None a string vacío)
        nombre_acudiente = str(aprendiz.acudiente_nombre_completo or "") if aprendiz else ""
        documento_acudiente = str(aprendiz.acudiente_documento or "") if aprendiz else ""
        tipo_documento_acudiente = str(aprendiz.acudiente_tipo_doc or "") if aprendiz else ""
        lugar_expedicion = str(aprendiz.acudiente_lugar_expedicion or "") if aprendiz else ""
        direccion_acudiente = str(aprendiz.acudiente_direccion or "") if aprendiz else ""
        telefono_acudiente = str(aprendiz.acudiente_telefono or "") if aprendiz else ""
        email_acudiente = str(aprendiz.acudiente_email or "") if aprendiz else ""

        # Datos académicos
        numero_grupo = str(aprendiz.grupo.nombre or "") if aprendiz and aprendiz.grupo else ""
        programa_nombre = str(aprendiz.programa.nombre or "ANÁLISIS Y DESARROLLO DE SOFTWARE") if aprendiz and aprendiz.programa else "ANÁLISIS Y DESARROLLO DE SOFTWARE"

        # Fecha actual
        fecha_actual = datetime.now().strftime('%d/%m/%Y')

        # Campos específicos del formato Tratamiento de Datos
        # Basados en los campos resaltados identificados
        replacements = {
            'Nombre Acudiente': nombre_acudiente,
            'Numero Documento': documento_acudiente.strip(),  # Eliminar espacios adicionales
            'Ciudad Documento': lugar_expedicion,
            'Nombre del aprendiz': nombre_aprendiz,
            'nombre del aprendiz': nombre_aprendiz,
            'Tipo de documento del aprendiz': tipo_documento_aprendiz,
            'número de documento del aprendiz': documento_aprendiz.strip(),
            'n°mero de documento del aprendiz': documento_aprendiz.strip(),
            'Numero Documento Aprendiz': documento_aprendiz.strip(),
            'Tipo Documento aprendiz, numero documento aprendiz': f'{tipo_documento_aprendiz} {documento_aprendiz.strip()}',
            'Tipo Documento acudiente, numero documento acuediente': f'{tipo_documento_acudiente} {documento_acudiente.strip()}',
            'Tipo Documento acudiente, numero documento acudiente': f'{tipo_documento_acudiente} {documento_acudiente.strip()}',
            'Fecha': fecha_actual,
            'Número Grupo': numero_grupo,
            'N°mero Grupo': numero_grupo,
            'Numero Grupo': numero_grupo,
            'numero grupo': numero_grupo,
            'numero de grupo': numero_grupo,
            'número de grupo': numero_grupo,
            'n°mero de grupo': numero_grupo,
            'Número de Grupo': numero_grupo,
            'N°mero de Grupo': numero_grupo,
            'Numero de Grupo': numero_grupo,
            'numero de ficha': numero_grupo,
            'número de ficha': numero_grupo,
            'Numero de Ficha': numero_grupo,
            'Número de Ficha': numero_grupo,
            'N°mero de Ficha': numero_grupo,
            'ficha': numero_grupo,
            'Ficha': numero_grupo,
            'correo acudiente': email_acudiente,
            'Correo acudiente': email_acudiente,
            'correo del acudiente': email_acudiente,
            'Correo del acudiente': email_acudiente,
            'dirección acudiente': direccion_acudiente,
            'direcci°n acudiente': direccion_acudiente,
            'Dirección acudiente': direccion_acudiente,
            'Direcci°n acudiente': direccion_acudiente,
            'dirección del acudiente': direccion_acudiente,
            'Dirección del acudiente': direccion_acudiente,
            'direcci°n del acudiente': direccion_acudiente,
            'Direcci°n del acudiente': direccion_acudiente,
            'Direccion acudiente': direccion_acudiente,
            'direccion acudiente': direccion_acudiente,
            'Direccion del acudiente': direccion_acudiente,
            'direccion del acudiente': direccion_acudiente,
            'Dirección de contacto: dirección acudiente': f'Dirección de contacto: {direccion_acudiente}',
            'Dirección de contacto: direcci°n acudiente': f'Dirección de contacto: {direccion_acudiente}',
            'Direcci°n de contacto: direcci°n acudiente': f'Dirección de contacto: {direccion_acudiente}',
            'Direcci�n de contacto: direcci�n acudiente': f'Dirección de contacto: {direccion_acudiente}',
            'Dirección de contacto: dirección acudiente': f'Dirección de contacto: {direccion_acudiente}',
            'direcci�n acudiente': direccion_acudiente,
            'Direcci�n acudiente': direccion_acudiente,
            'Teléfono acudiente': telefono_acudiente,
            'Telefono acudiente': telefono_acudiente,
            'teléfono acudiente': telefono_acudiente,
            'telefono acudiente': telefono_acudiente,
            'Teléfono del acudiente': telefono_acudiente,
            'Telefono del acudiente': telefono_acudiente,
            'teléfono del acudiente': telefono_acudiente,
            'telefono del acudiente': telefono_acudiente,
            'ANÁLISIS Y DESARROLLO DE SOFTWARE ': programa_nombre,
            'AN…LISIS Y DESARROLLO DE SOFTWARE ': programa_nombre,
            'ANALISIS Y DESARROLLO DE SOFTWARE': programa_nombre,
            'Análisis y Desarrollo de Software': programa_nombre,
        }

        logger.info(f"Generando formato de tratamiento de datos para {nombre_aprendiz}")

        # Procesar documento en UNA SOLA PASADA (optimizado)
        checkbox_procesado = False

        # Procesar párrafos (checkboxes primero, luego reemplazos + highlights en una pasada)
        for paragraph in doc.paragraphs:
            # 1. Marcar checkboxes PRIMERO (antes de modificar runs con reemplazos)
            if not checkbox_procesado:
                text = paragraph.text
                if 'Ciudadan' in text and 'Extranjer' in text:
                    x_runs = [i for i, run in enumerate(paragraph.runs) if run.text.strip().lower() == 'x']

                    if len(x_runs) == 2:
                        cc_run_idx, ce_run_idx = x_runs[0], x_runs[1]

                        if tipo_documento_acudiente == 'CC':
                            paragraph.runs[cc_run_idx].text = 'X'
                            paragraph.runs[ce_run_idx].text = ' '
                        elif tipo_documento_acudiente in ['CE', 'PPT', 'PEP']:
                            paragraph.runs[cc_run_idx].text = ' '
                            paragraph.runs[ce_run_idx].text = 'X'
                        else:
                            paragraph.runs[cc_run_idx].text = ' '
                            paragraph.runs[ce_run_idx].text = ' '

                        checkbox_procesado = True

            # 2. Reemplazar texto (después de checkboxes)
            replace_text_in_paragraph(paragraph, replacements)

            # 3. Eliminar highlights
            for run in paragraph.runs:
                run.font.highlight_color = None

        # Procesar tablas (reemplazos + highlights en una pasada)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                        # Eliminar highlights
                        for run in paragraph.runs:
                            run.font.highlight_color = None

        # Guardar el documento generado
        # Obtener la ruta base del proyecto (ir dos niveles arriba desde este archivo)
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        temp_dir = os.path.join(base_dir, 'temp')

        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        output_filename = f"Tratamiento_Datos_{documento_aprendiz}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(temp_dir, output_filename)

        doc.save(output_path)

        # Convertir a PDF
        pdf_path = convert_docx_to_pdf(output_path)
        return pdf_path

    except Exception as e:
        raise Exception(f"Error al generar formato de tratamiento de datos: {str(e)}")


def generar_formato_compromiso_aprendiz(aprendiz):
    """
    Genera el formato de Compromiso del Aprendiz con los campos resaltados autocompletados

    Args:
        aprendiz: Objeto Aprendiz con los datos del estudiante y acudiente

    Returns:
        str: Ruta del archivo generado
    """
    try:
        # Cargar la plantilla base
        template_path = os.path.join('formatos', 'Compromiso del Aprendiz.docx')
        doc = Document(template_path)

        # Datos del aprendiz (convertir None a string vacío)
        nombre_aprendiz = str(aprendiz.usuario.nombre_completo or "") if aprendiz and aprendiz.usuario else ""
        documento_aprendiz = str(aprendiz.usuario.documento or "") if aprendiz and aprendiz.usuario else ""
        tipo_documento_aprendiz = str(aprendiz.usuario.tipo_documento or "") if aprendiz and aprendiz.usuario else ""

        # Datos del programa (convertir None a string vacío)
        programa_nombre = str(aprendiz.programa.nombre or "ANÁLISIS Y DESARROLLO DE SOFTWARE") if aprendiz and aprendiz.programa else "ANÁLISIS Y DESARROLLO DE SOFTWARE"

        # Datos del grupo (convertir None a string vacío)
        grupo_nombre = str(aprendiz.grupo.nombre or "") if aprendiz and aprendiz.grupo else ""

        # Datos del acudiente (están en el modelo Aprendiz - convertir None a string vacío)
        documento_acudiente = str(aprendiz.acudiente_documento or "") if aprendiz else ""
        tipo_documento_acudiente = str(aprendiz.acudiente_tipo_doc or "") if aprendiz else ""

        # Fecha actual
        fecha_actual = datetime.now()
        dia = fecha_actual.strftime('%d')
        mes = fecha_actual.strftime('%m')
        ano = fecha_actual.strftime('%Y')

        # Campos específicos del formato Compromiso del Aprendiz
        # Basados en los campos resaltados identificados
        replacements = {
            'Nombre aprendiz': nombre_aprendiz,
            'Nombre\n aprendiz': nombre_aprendiz,
            ' Nombre aprendiz': nombre_aprendiz,
            'nombre aprendiz': nombre_aprendiz,
            'nombre del aprendiz': nombre_aprendiz,
            'Nombre del aprendiz': nombre_aprendiz,
            'Número Documento Aprendiz': documento_aprendiz,
            'N°mero Documento Aprendiz': documento_aprendiz,
            'Numero Documento aprendiz': documento_aprendiz,
            'número documento aprendiz': documento_aprendiz,
            'numero documento aprendiz': documento_aprendiz,
            # Programa - múltiples variaciones
            # El template tiene "Matriculado en el programa de formación: nombre del programa o tecnica"
            # Solo necesitamos reemplazar "nombre del programa o tecnica" por el nombre real
            'nombre del programa o tecnica': programa_nombre,
            'Nombre del programa o tecnica': programa_nombre,
            'nombre del programa o técnica': programa_nombre,
            'Nombre del programa o técnica': programa_nombre,
            'NOMBRE DEL PROGRAMA O TECNICA': programa_nombre,
            'NOMBRE DEL PROGRAMA O TÉCNICA': programa_nombre,
            'nombre del programa': programa_nombre,
            'Nombre del programa': programa_nombre,
            'NOMBRE DEL PROGRAMA': programa_nombre,
            'nombre programa': programa_nombre,
            'Nombre programa': programa_nombre,
            'NOMBRE PROGRAMA': programa_nombre,
            # Número de ficha/grupo - múltiples variaciones
            # Nota: "Ficha de Caracterización No." se mantiene como está en la plantilla
            'Numero de Grupo': grupo_nombre,
            'N°mero de Grupo': grupo_nombre,
            'Número de Grupo': grupo_nombre,
            'numero de grupo': grupo_nombre,
            'número de grupo': grupo_nombre,
            'Numero Grupo': grupo_nombre,
            'Número Grupo': grupo_nombre,
            'N°mero Grupo': grupo_nombre,
            'numero grupo': grupo_nombre,
            'número grupo': grupo_nombre,
            'ficha de caracterización': grupo_nombre,
            'Ficha de caracterización': grupo_nombre,
            'ficha de caracterizacion': grupo_nombre,
            'Ficha de caracterizacion': grupo_nombre,
            'numero de ficha': grupo_nombre,
            'número de ficha': grupo_nombre,
            'Numero de ficha': grupo_nombre,
            'Número de ficha': grupo_nombre,
            'N°mero de ficha': grupo_nombre,
            'numero ficha': grupo_nombre,
            'número ficha': grupo_nombre,
            'Numero ficha': grupo_nombre,
            'Número ficha': grupo_nombre,
            # Nota: NO reemplazar "ficha" o "Ficha" genéricamente porque afecta "Ficha de Caracterización No."
            # Datos del acudiente
            'Tipo documento acudiente, numero documento acudiente': f'{tipo_documento_acudiente} {documento_acudiente}',
            # Fechas
            'dia de hoy': dia,
            'dia': dia,
            'día': dia,
            'Dia': dia,
            'Día': dia,
            'mes actual': mes,
            'Mes actual': mes,
            'mes': mes,
            'Mes': mes,
            'año actual completo': ano,
            'a°o actual completo': ano,
            'año actual': ano,
            'a°o actual': ano,
            'año': ano,
            'a°o': ano,
            'Año': ano,
            'A°o': ano,
        }

        logger.info(f"Generando formato de compromiso del aprendiz para {nombre_aprendiz}")
        # Procesar documento en UNA SOLA PASADA (optimizado)
        # Procesar párrafos (reemplazos + highlights)
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
            # Eliminar highlights
            for run in paragraph.runs:
                run.font.highlight_color = None

        # Procesar tablas (reemplazos + checkboxes + highlights en una pasada)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 1. Reemplazar texto
                        replace_text_in_paragraph(paragraph, replacements)

                        # 2. Marcar checkboxes
                        for run in paragraph.runs:
                            text = run.text.strip()
                            # Procesar tipos de documento
                            if text in ['TI', 'CC', 'PPT', 'CE', 'PEP']:
                                if text == tipo_documento_aprendiz:
                                    run.text = 'X'
                                elif text == 'PPT' and tipo_documento_aprendiz in ['PPT', 'CE', 'PEP']:
                                    run.text = 'X'
                                else:
                                    run.text = ' '

                            # 3. Eliminar highlights
                            run.font.highlight_color = None

        # Guardar el documento generado
        # Obtener la ruta base del proyecto (ir dos niveles arriba desde este archivo)
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        temp_dir = os.path.join(base_dir, 'temp')

        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        output_filename = f"Compromiso_Aprendiz_{documento_aprendiz}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(temp_dir, output_filename)

        doc.save(output_path)

        # Convertir a PDF
        pdf_path = convert_docx_to_pdf(output_path)
        return pdf_path

    except Exception as e:
        raise Exception(f"Error al generar formato de compromiso del aprendiz: {str(e)}")


def generar_pdf_unificado_aprendiz(aprendiz, documentos):
    """
    Genera un PDF unificado con todos los documentos del aprendiz

    Args:
        aprendiz: Objeto Aprendiz con los datos del estudiante
        documentos: Lista de objetos Documento con los archivos del aprendiz

    Returns:
        str: Ruta del archivo PDF generado
    """
    try:
        if not documentos:
            raise Exception("No hay documentos para unificar")

        # Crear carpeta temp si no existe
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        temp_dir = os.path.join(base_dir, 'temp')

        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        # Nombre del archivo PDF unificado
        nombre_aprendiz = aprendiz.usuario.nombre_completo.replace(' ', '_')
        documento_aprendiz = aprendiz.usuario.documento
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"Documentos_{nombre_aprendiz}_{documento_aprendiz}_{timestamp}.pdf"
        output_path = os.path.join(temp_dir, output_filename)

        # Crear el merger de PDFs
        merger = PdfMerger()

        # Procesar cada documento
        for documento in documentos:
            file_path = documento.ruta_archivo

            # Convertir ruta relativa a absoluta si es necesario
            if not os.path.isabs(file_path):
                base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                file_path = os.path.join(base_dir, file_path)

            if os.path.exists(file_path):
                # Solo agregar archivos PDF
                if file_path.lower().endswith('.pdf'):
                    try:
                        merger.append(file_path)
                        logger.info(f"Agregado: {documento.tipo_documento} - {documento.nombre_archivo}")
                    except Exception as e:
                        logger.warning(f"No se pudo agregar {documento.nombre_archivo}: {str(e)}")
                else:
                    logger.warning(f"Archivo no es PDF, se omitió: {documento.nombre_archivo}")
            else:
                logger.warning(f"Archivo no encontrado: {file_path}")

        # Guardar el PDF unificado
        merger.write(output_path)
        merger.close()

        logger.info(f"PDF unificado generado exitosamente: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Error al generar PDF unificado: {str(e)}")
        raise Exception(f"Error al generar PDF unificado: {str(e)}")


def generar_pdf_unificado_grupo(grupo, aprendices_lista):
    """
    Genera un PDF unificado con todos los documentos de un grupo

    Args:
        grupo: Objeto Grupo con los datos del grupo
        aprendices_lista: Lista de objetos Aprendiz del grupo

    Returns:
        str: Ruta del archivo PDF generado
    """
    from app.models.matricula import Matricula
    from app.models.documento import Documento

    try:
        if not aprendices_lista:
            raise Exception("No hay aprendices en el grupo")

        # Crear carpeta temp si no existe
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        temp_dir = os.path.join(base_dir, 'temp')

        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        # Nombre del archivo PDF unificado
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"Documentos_Grupo_{grupo.nombre}_{timestamp}.pdf"
        output_path = os.path.join(temp_dir, output_filename)

        # Crear el merger de PDFs
        merger = PdfMerger()
        documentos_agregados = 0

        # Procesar cada aprendiz
        for aprendiz in aprendices_lista:
            # Obtener la matrícula del aprendiz
            matricula = Matricula.query.filter_by(aprendiz_id=aprendiz.id).first()
            if not matricula:
                continue

            # Obtener documentos activos
            documentos = Documento.query.filter_by(matricula_id=matricula.id).filter(
                Documento.reemplazado_por == None
            ).all()

            if not documentos:
                continue

            # Procesar cada documento del aprendiz
            for documento in documentos:
                file_path = documento.ruta_archivo

                # Convertir ruta relativa a absoluta si es necesario
                if not os.path.isabs(file_path):
                    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                    file_path = os.path.join(base_dir, file_path)

                if os.path.exists(file_path):
                    # Solo agregar archivos PDF
                    if file_path.lower().endswith('.pdf'):
                        try:
                            merger.append(file_path)
                            documentos_agregados += 1
                            logger.info(f"Agregado: {aprendiz.usuario.nombre_completo} - {documento.tipo_documento}")
                        except Exception as e:
                            logger.warning(f"No se pudo agregar {documento.nombre_archivo}: {str(e)}")
                    else:
                        logger.warning(f"Archivo no es PDF, se omitió: {documento.nombre_archivo}")
                else:
                    logger.warning(f"Archivo no encontrado: {file_path}")

        if documentos_agregados == 0:
            raise Exception("No se encontraron documentos PDF para unificar")

        # Guardar el PDF unificado
        merger.write(output_path)
        merger.close()

        logger.info(f"PDF unificado del grupo generado exitosamente: {output_path} ({documentos_agregados} documentos)")
        return output_path

    except Exception as e:
        logger.error(f"Error al generar PDF unificado del grupo: {str(e)}")
        raise Exception(f"Error al generar PDF unificado del grupo: {str(e)}")
